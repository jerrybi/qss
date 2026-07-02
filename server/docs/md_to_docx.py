import re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F6F8FA'
    })
    pPr.append(shd)
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    for side in ['top', 'left', 'bottom', 'right']:
        border = pBdr.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '4', qn('w:color'): 'E1E4E8'
        })
        pBdr.append(border)
    pPr.append(pBdr)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2F)

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r') as f:
        lines = f.readlines()

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        if line.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_lines))
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == '':
            i += 1
            continue

        if line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            text = line[level:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_heading(text, level=min(level, 6))
            i += 1
            continue

        if line.startswith('>'):
            text = line.lstrip('>').strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            left = pBdr.makeelement(qn('w:left'), {
                qn('w:val'): 'single', qn('w:sz'): '18',
                qn('w:space'): '8', qn('w:color'): '4CAF50'
            })
            pBdr.append(left)
            pPr.append(pBdr)
            i += 1
            continue

        # Table detection
        if '|' in line and i + 1 < len(lines) and '---' in lines[i+1]:
            headers = [h.strip() for h in line.split('|') if h.strip()]
            i += 2
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                row_data = [c.strip() for c in lines[i].split('|') if c.strip()]
                rows.append(row_data)
                i += 1

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Light Shading Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', h)
                clean = re.sub(r'`(.*?)`', r'\1', clean)
                cell.text = clean
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

            for r_idx, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < len(headers):
                        cell = table.rows[r_idx + 1].cells[j]
                        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                        clean = re.sub(r'`(.*?)`', r'\1', clean)
                        clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean)
                        cell.text = clean
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

            doc.add_paragraph()
            continue

        # Bullet list
        if re.match(r'^[\-\*]\s', line):
            text = re.sub(r'^[\-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                else:
                    p.add_run(part)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\[.*?\]\(.*?\))', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                elif part.startswith('[') and '](' in part:
                    m = re.match(r'\[(.*?)\]\(', part)
                    if m:
                        run = p.add_run(m.group(1))
                        run.font.color.rgb = RGBColor(0x03, 0x66, 0xD6)
                        run.font.underline = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\[.*?\]\(.*?\))', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            elif part.startswith('[') and '](' in part:
                m = re.match(r'\[(.*?)\]\(', part)
                if m:
                    run = p.add_run(m.group(1))
                    run.font.color.rgb = RGBColor(0x03, 0x66, 0xD6)
                    run.font.underline = True
            else:
                p.add_run(part)
        i += 1

    if in_code_block and code_lines:
        add_code_block(doc, '\n'.join(code_lines))

    doc.save(docx_path)
    return docx_path


admin = md_to_docx(
    '/Volumes/projects/qss/project/server/docs/api-key-admin-guide.md',
    '/Volumes/projects/qss/project/server/docs/api-key-admin-guide.docx'
)
print(f"Generated: {admin}")

dev = md_to_docx(
    '/Volumes/projects/qss/project/server/docs/api-v1-exhibitor-guide.md',
    '/Volumes/projects/qss/project/server/docs/api-v1-exhibitor-guide.docx'
)
print(f"Generated: {dev}")
